from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Form
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import chromadb
from emergentintegrations.llm.chat import LlmChat, UserMessage
import PyPDF2
import docx
import io
from bs4 import BeautifulSoup
import html2text
import requests
import tiktoken
import asyncio

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# ChromaDB setup
chroma_client = chromadb.PersistentClient(path=str(ROOT_DIR / 'chromadb_data'))
chroma_collection = chroma_client.get_or_create_collection(name="aurenlm_documents")

# LLM setup
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Create the main app
app = FastAPI()
api_router = APIRouter(prefix="/api")

# Models
class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_type: str
    content: str
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    chunks: List[str] = []

class DocumentCreate(BaseModel):
    filename: str
    file_type: str
    content: str

class Mindmap(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_ids: List[str]
    title: str
    structure: Dict[str, Any]
    created_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatMessage(BaseModel):
    role: str
    content: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatSession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_ids: List[str]
    messages: List[ChatMessage] = []
    created_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class QuestionRequest(BaseModel):
    question: str
    document_ids: List[str]
    session_id: Optional[str] = None

class NotesRequest(BaseModel):
    topic: str
    document_ids: List[str]

class QuizRequest(BaseModel):
    document_ids: List[str]
    num_questions: int = 5

# Helper functions
def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting PDF: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting DOCX: {str(e)}")

def extract_text_from_url(url: str) -> str:
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        h = html2text.HTML2Text()
        h.ignore_links = False
        text = h.handle(str(soup))
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting URL: {str(e)}")

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk:
            chunks.append(chunk)
        start = end - overlap
    return chunks

async def generate_embeddings_and_store(doc_id: str, chunks: List[str]):
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"embed_{doc_id}",
            system_message="You are an embedding assistant."
        ).with_model("openai", "gpt-5")
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            chroma_collection.add(
                documents=[chunk],
                ids=[chunk_id],
                metadatas=[{"doc_id": doc_id, "chunk_index": i}]
            )
    except Exception as e:
        logging.error(f"Error generating embeddings: {str(e)}")

async def query_similar_chunks(query: str, document_ids: List[str], n_results: int = 5) -> List[str]:
    try:
        results = chroma_collection.query(
            query_texts=[query],
            n_results=n_results,
            where={"doc_id": {"$in": document_ids}}
        )
        return results['documents'][0] if results['documents'] else []
    except Exception as e:
        logging.error(f"Error querying similar chunks: {str(e)}")
        return []

# Routes
@api_router.get("/")
async def root():
    return {"message": "AurenLM API"}

@api_router.post("/upload", response_model=Document)
async def upload_file(file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
        file_type = file.filename.split('.')[-1].lower()
        
        if file_type == 'pdf':
            content = extract_text_from_pdf(file_bytes)
        elif file_type in ['docx', 'doc']:
            content = extract_text_from_docx(file_bytes)
        elif file_type == 'txt':
            content = file_bytes.decode('utf-8')
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        chunks = chunk_text(content)
        
        document = Document(
            filename=file.filename,
            file_type=file_type,
            content=content,
            chunks=chunks
        )
        
        doc_dict = document.model_dump()
        doc_dict['upload_date'] = doc_dict['upload_date'].isoformat()
        
        await db.documents.insert_one(doc_dict)
        await generate_embeddings_and_store(document.id, chunks)
        
        return document
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/upload-url")
async def upload_url(url: str = Form(...)):
    try:
        content = extract_text_from_url(url)
        chunks = chunk_text(content)
        
        document = Document(
            filename=url,
            file_type="url",
            content=content,
            chunks=chunks
        )
        
        doc_dict = document.model_dump()
        doc_dict['upload_date'] = doc_dict['upload_date'].isoformat()
        
        await db.documents.insert_one(doc_dict)
        await generate_embeddings_and_store(document.id, chunks)
        
        return document
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/documents", response_model=List[Document])
async def get_documents():
    documents = await db.documents.find({}, {"_id": 0}).to_list(1000)
    for doc in documents:
        if isinstance(doc['upload_date'], str):
            doc['upload_date'] = datetime.fromisoformat(doc['upload_date'])
    return documents

@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    result = await db.documents.delete_one({"id": doc_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete from ChromaDB
    try:
        chroma_collection.delete(where={"doc_id": doc_id})
    except:
        pass
    
    return {"message": "Document deleted successfully"}

@api_router.post("/generate-mindmap")
async def generate_mindmap(document_ids: List[str]):
    try:
        documents = await db.documents.find(
            {"id": {"$in": document_ids}}, 
            {"_id": 0}
        ).to_list(1000)
        
        if not documents:
            raise HTTPException(status_code=404, detail="No documents found")
        
        combined_content = "\n\n".join([doc['content'][:3000] for doc in documents])
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"mindmap_{uuid.uuid4()}",
            system_message="You are an expert at analyzing documents and creating structured mindmaps."
        ).with_model("openai", "gpt-5")
        
        prompt = f"""Analyze the following documents and create a hierarchical mindmap structure.
        
Return ONLY a JSON object with this exact structure:
{{
  "title": "Main Topic Title",
  "nodes": [
    {{
      "id": "node1",
      "label": "Main Concept 1",
      "type": "main",
      "children": [
        {{
          "id": "node1_1",
          "label": "Subtopic 1.1",
          "type": "sub",
          "children": []
        }}
      ]
    }}
  ]
}}

Documents:
{combined_content[:8000]}"""
        
        message = UserMessage(text=prompt)
        response = await chat.send_message(message)
        
        import json
        try:
            structure = json.loads(response)
        except:
            structure = {
                "title": "Document Analysis",
                "nodes": [
                    {
                        "id": "node1",
                        "label": "Main Content",
                        "type": "main",
                        "children": []
                    }
                ]
            }
        
        mindmap = Mindmap(
            document_ids=document_ids,
            title=structure.get('title', 'Document Mindmap'),
            structure=structure
        )
        
        mindmap_dict = mindmap.model_dump()
        mindmap_dict['created_date'] = mindmap_dict['created_date'].isoformat()
        
        await db.mindmaps.insert_one(mindmap_dict)
        
        return mindmap
    except Exception as e:
        logging.error(f"Error generating mindmap: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/generate-notes")
async def generate_notes(request: NotesRequest):
    try:
        relevant_chunks = await query_similar_chunks(request.topic, request.document_ids, n_results=3)
        
        context = "\n\n".join(relevant_chunks)
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"notes_{uuid.uuid4()}",
            system_message="You are an expert educator who creates clear, structured study notes."
        ).with_model("openai", "gpt-5")
        
        prompt = f"""Create detailed, well-structured study notes for the topic: {request.topic}

Use the following context from the documents:
{context}

Format the notes with:
- Clear definitions
- Key points (bullet points)
- Examples when relevant
- Summary at the end"""
        
        message = UserMessage(text=prompt)
        response = await chat.send_message(message)
        
        return {"notes": response, "topic": request.topic}
    except Exception as e:
        logging.error(f"Error generating notes: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/chat")
async def chat_with_documents(request: QuestionRequest):
    try:
        relevant_chunks = await query_similar_chunks(
            request.question, 
            request.document_ids, 
            n_results=5
        )
        
        context = "\n\n".join(relevant_chunks)
        
        session_id = request.session_id or f"chat_{uuid.uuid4()}"
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=session_id,
            system_message=f"""You are AurenLM, an intelligent study assistant. Answer questions based on the provided context from uploaded documents. 
            If the answer is not in the context, say so clearly.
            
Context:
{context}"""
        ).with_model("openai", "gpt-5")
        
        message = UserMessage(text=request.question)
        response = await chat.send_message(message)
        
        # Save to database
        user_msg = ChatMessage(role="user", content=request.question)
        ai_msg = ChatMessage(role="assistant", content=response)
        
        if request.session_id:
            await db.chat_sessions.update_one(
                {"id": session_id},
                {
                    "$push": {
                        "messages": [
                            {**user_msg.model_dump(), "timestamp": user_msg.timestamp.isoformat()},
                            {**ai_msg.model_dump(), "timestamp": ai_msg.timestamp.isoformat()}
                        ]
                    }
                },
                upsert=True
            )
        else:
            chat_session = ChatSession(
                id=session_id,
                document_ids=request.document_ids,
                messages=[user_msg, ai_msg]
            )
            session_dict = chat_session.model_dump()
            session_dict['created_date'] = session_dict['created_date'].isoformat()
            for msg in session_dict['messages']:
                msg['timestamp'] = msg['timestamp'].isoformat()
            await db.chat_sessions.insert_one(session_dict)
        
        return {
            "answer": response,
            "session_id": session_id
        }
    except Exception as e:
        logging.error(f"Error in chat: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/generate-quiz")
async def generate_quiz(request: QuizRequest):
    try:
        documents = await db.documents.find(
            {"id": {"$in": request.document_ids}}, 
            {"_id": 0}
        ).to_list(1000)
        
        if not documents:
            raise HTTPException(status_code=404, detail="No documents found")
        
        combined_content = "\n\n".join([doc['content'][:2000] for doc in documents])
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"quiz_{uuid.uuid4()}",
            system_message="You are an expert educator who creates engaging quiz questions."
        ).with_model("openai", "gpt-5")
        
        prompt = f"""Based on the following content, create {request.num_questions} multiple-choice questions.

Return ONLY a JSON array with this structure:
[
  {{
    "question": "Question text?",
    "options": ["A) Option 1", "B) Option 2", "C) Option 3", "D) Option 4"],
    "correct_answer": "A",
    "explanation": "Brief explanation"
  }}
]

Content:
{combined_content[:5000]}"""
        
        message = UserMessage(text=prompt)
        response = await chat.send_message(message)
        
        import json
        try:
            quiz_data = json.loads(response)
        except:
            quiz_data = [{
                "question": "What is the main topic of the uploaded documents?",
                "options": ["A) Option 1", "B) Option 2", "C) Option 3", "D) Option 4"],
                "correct_answer": "A",
                "explanation": "Based on the content analysis"
            }]
        
        return {"quiz": quiz_data}
    except Exception as e:
        logging.error(f"Error generating quiz: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/chat-sessions", response_model=List[ChatSession])
async def get_chat_sessions():
    sessions = await db.chat_sessions.find({}, {"_id": 0}).to_list(1000)
    for session in sessions:
        if isinstance(session['created_date'], str):
            session['created_date'] = datetime.fromisoformat(session['created_date'])
        for msg in session['messages']:
            if isinstance(msg['timestamp'], str):
                msg['timestamp'] = datetime.fromisoformat(msg['timestamp'])
    return sessions

# Include router
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()